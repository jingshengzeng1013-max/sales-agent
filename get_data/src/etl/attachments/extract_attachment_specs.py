# -*- coding: utf-8 -*-
"""
附件技术参数提取模块
使用本地 LLM 直接解析 PDF 文件，提取技术参数和采购产品表格
输出 JSONL 供向量检索用
"""

import os
import sys
import json
import time
import re
import base64
from pathlib import Path

# 添加项目根目录到路径
BASE_DIR = Path(__file__).resolve().parent.parent.parent.parent
sys.path.append(str(BASE_DIR))

from src.config import LOCAL_LLM_CONFIG, DB_PATH
from src.utils.jsonl_helper import load_jsonl, save_jsonl

# 尝试导入依赖
try:
    from openai import OpenAI
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False
    print("[ERROR] openai 库未安装，请运行：pip install openai")

try:
    import fitz  # PyMuPDF
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False
    print("[WARN] PyMuPDF 未安装，PDF 图片解析不可用")

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import olefile
    HAS_OLEFILE = True
except ImportError:
    HAS_OLEFILE = False

# 附件目录
ATTACHMENTS_DIR = BASE_DIR / "data" / "attachment"


def get_local_llm_client():
    """获取本地 LLM 客户端"""
    if not HAS_OPENAI:
        return None
    return OpenAI(
        api_key=LOCAL_LLM_CONFIG.get("api_key", "local-api-key"),
        base_url=LOCAL_LLM_CONFIG["base_url"],
        timeout=LOCAL_LLM_CONFIG.get("timeout", 600)
    )


def pdf_to_images(pdf_path: str, start_page: int = 0, end_page: int = None, dpi: int = 150) -> list:
    """
    将 PDF 页面转换为图片

    Args:
        pdf_path: PDF 文件路径
        start_page: 起始页 (0索引)
        end_page: 结束页 (不包含)，None 表示最后一页
        dpi: 图片清晰度

    Returns:
        list: 图片 base64 列表
    """
    if not HAS_FITZ:
        return []

    images = []
    try:
        doc = fitz.open(pdf_path)
        end_page = end_page or doc.page_count

        for page_num in range(start_page, min(end_page, doc.page_count)):
            page = doc[page_num]
            # 缩放因子，控制输出图片大小
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes("png")
            img_b64 = base64.b64encode(img_bytes).decode("utf-8")
            images.append(img_b64)

        doc.close()
    except Exception as e:
        print(f"  [ERROR] PDF 转图片失败: {e}")

    return images


def find_table_pages(pdf_path: str, keywords: list = None) -> list:
    """
    搜索 PDF 中可能包含表格的页面

    Args:
        pdf_path: PDF 文件路径
        keywords: 关键词列表

    Returns:
        list: 可能包含表格的页码列表 (0索引)
    """
    if not HAS_FITZ or keywords is None:
        return list(range(20))  # 默认前20页

    keywords = keywords or ["序号", "产品", "规格", "技术参数", "数量", "单位", "单价", "合计", "预算", "采购"]

    table_pages = []
    try:
        doc = fitz.open(pdf_path)
        for i in range(doc.page_count):
            text = doc[i].get_text()
            # 统计关键词匹配数量
            match_count = sum(1 for kw in keywords if kw in text)
            if match_count >= 2:  # 至少2个关键词匹配
                table_pages.append(i)
            if len(table_pages) >= 10:  # 最多取10页
                break
        doc.close()
    except Exception as e:
        print(f"  [WARN] 搜索表格页失败: {e}")

    return table_pages if table_pages else list(range(min(5, doc.page_count)))


def extract_tables_with_vision(images: list, file_name: str = "", tender_id: str = "") -> dict:
    """
    使用视觉模型从 PDF 图片中提取表格数据

    Args:
        images: PDF 页面图片 base64 列表
        file_name: 文件名
        tender_id: 招标项目 ID

    Returns:
        dict: 提取的表格数据
    """
    client = get_local_llm_client()
    if not client:
        return None

    if not images:
        return None

    # 构建图片消息
    image_contents = []
    for i, img_b64 in enumerate(images):
        image_contents.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/png;base64,{img_b64}"}
        })

    prompt = """你是一个专业的招投标数据抽取专家。请从招标文档图片中提取采购产品清单表格。

## 必须提取的字段
请提取"货物需求明细"或"采购清单"表格中的每一项：
- 序号
- 产品名称/品目
- 规格型号/技术参数
- 数量
- 单位
- 单价(元)（如有）
- 合计(元)（如有）

## 输出要求
1. 输出严格 JSON 格式，不要有任何额外文字
2. 所有数据必须来源于图片，不要猜测或编造
3. 如果某字段为空，设为空字符串
4. items 是数组，每行产品一个对象

## 重要：格式要求
每一项用【项目名：技术需求】的格式呈现：
- 【产品名称】灭火救援装备
- 【规格型号】XXX型号
- 【技术参数】详细技术要求描述
- 【数量】10
- 【单位】套
- 【单价】5000
- 【合计】50000

## 输出 JSON 格式
{
  "items": [
    {
      "序号": "1",
      "产品名称": "灭火救援装备",
      "规格型号": "XXX型号",
      "技术参数": "详细技术要求描述",
      "数量": "10",
      "单位": "套",
      "单价": "5000",
      "合计": "50000"
    }
  ],
  "total_budget": "总金额",
  "project_name": "项目名称(如果可见)",
  "buyer_name": "采购单位(如果可见)"
}"""

    messages = [{
        "role": "user",
        "content": [
            {"type": "text", "text": prompt},
            *image_contents
        ]
    }]

    try:
        response = client.chat.completions.create(
            model=LOCAL_LLM_CONFIG["model"],
            messages=messages,
            temperature=0.1,
            max_tokens=4000
        )

        result = response.choices[0].message.content

        # 提取 JSON
        json_match = re.search(r'\{.*\}', result, re.DOTALL)
        if json_match:
            raw_json = json_match.group(0)

            # 如果完整 JSON 解析失败，尝试手动解析 items
            try:
                data = json.loads(raw_json)
            except json.JSONDecodeError as e:
                print(f"  [WARN] JSON 解析失败: {e}")
                # 尝试提取 items 数组部分
                items_match = re.search(r'"items"\s*:\s*\[(.*)\]', raw_json, re.DOTALL)
                if items_match:
                    items_str = items_match.group(1)
                    items = []
                    # 尝试逐个提取 item 对象
                    item_matches = re.finditer(r'\{[^{}]*\}', items_str)
                    for im in item_matches:
                        try:
                            item = json.loads(im.group(0))
                            items.append(item)
                        except:
                            continue
                    data = {"items": items}
                else:
                    data = {"items": []}

            data["file_name"] = file_name
            data["tender_id"] = tender_id
            data["source_type"] = "pdf_vision"

            # 将 items 转换为 【项目名：技术需求】 格式
            items = data.get("items", [])
            formatted_items = []
            for item in items:
                if not isinstance(item, dict):
                    continue
                product_name = item.get("产品名称", "")
                # 组合所有技术参数为一个技术需求字符串
                tech_params = []
                for key in ["规格型号", "技术参数", "数量", "单位", "单价", "合计"]:
                    val = item.get(key, "")
                    if val:
                        tech_params.append(f"{key}: {val}")
                tech_requirement = "; ".join(tech_params)

                formatted_items.append({
                    "产品名称": product_name,
                    "技术需求": tech_requirement
                })

            data["formatted_items"] = formatted_items
        else:
            print(f"  [WARN] 未找到 JSON: {result[:200] if result else 'None'}")
            return None

    except Exception as e:
        print(f"  [ERROR] LLM 调用失败: {e}")
        return None


def extract_text_with_llm(text_content: str, file_name: str = "", tender_id: str = "") -> dict:
    """
    使用文本模型从附件文本中提取技术参数（PDF 或其他格式）

    Args:
        text_content: 附件原始文本
        file_name: 原始文件名
        tender_id: 招标项目 ID

    Returns:
        dict: 提取的结构化技术参数
    """
    client = get_local_llm_client()
    if not client:
        return None

    # 截断过长文本
    max_chars = 15000
    if len(text_content) > max_chars:
        text_content = text_content[:max_chars] + "\n...（内容截断）"

    prompt = f"""你是一个专业的招投标技术参数抽取专家。请从以下附件文档中提取所有技术规格和参数。

## 文件信息
- 文件名：{file_name}
- 招标项目ID：{tender_id}

## 任务
从文档中提取：
1. **产品名称/品目**：采购的具体产品名称
2. **规格型号**：具体的技术规格、型号、尺寸、材质等
3. **技术参数**：数量、单位、性能指标、功能要求
4. **配置清单**：设备清单、软件清单、配件清单
5. **质量标准**：执行标准、认证要求
6. **交货要求**：交货期、验收标准
7. **技术规格原文**：文档中原文描述的重要技术段落

## 输出要求
1. 输出严格 JSON 格式，不要有任何额外文字
2. 所有字段必须来源于文档原文，不要猜测
3. 如果某项不存在，设为空数组或空字符串
4. 技术规格原文要保留原文的关键描述

## 输出 JSON 格式
{{
  "product_names": ["产品名称1", "产品名称2"],
  "specifications": ["规格型号1", "规格型号2"],
  "technical_params": ["参数1: 值1", "参数2: 值2"],
  "quantity_unit": "数量单位描述",
  "config_list": ["配置项1", "配置项2"],
  "quality_standards": ["执行标准1", "标准2"],
  "delivery_requirements": "交货和验收要求",
  "technical_paragraphs": ["原文技术描述段落1", "段落2"],
  "raw_snippets": ["包含关键参数的原文片段"]
}}

## 附件文档内容
{text_content}

请提取技术参数并返回 JSON："""

    try:
        response = client.chat.completions.create(
            model=LOCAL_LLM_CONFIG["model"],
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=3000
        )

        result = response.choices[0].message.content

        # 提取 JSON
        json_match = re.search(r'\{.*\}', result, re.DOTALL)
        if json_match:
            data = json.loads(json_match.group(0))
            data["file_name"] = file_name
            data["tender_id"] = tender_id
            data["source_text_length"] = len(text_content)
            data["source_type"] = "text"
            return data
        else:
            print(f"  [WARN] 未找到 JSON: {result[:200]}")
            return None

    except Exception as e:
        print(f"  [ERROR] LLM 调用失败: {e}")
        return None


def read_docx(file_path: str) -> str:
    """读取 DOCX 文件内容"""
    if not HAS_DOCX:
        return None
    try:
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                text_parts.append(para.text)
        return " ".join(text_parts)
    except Exception as e:
        print(f"  [ERROR] DOCX 读取失败: {e}")
        return None


def read_doc(file_path: str) -> str:
    """读取老版本 DOC 文件内容（使用 olefile）"""
    if not HAS_OLEFILE:
        # 尝试用 python-docx 读取
        try:
            from docx import Document as DocxDocument
            # DOC 实际上也可能被 python-docx 打开
            doc = DocxDocument(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    text_parts.append(para.text)
            return " ".join(text_parts)
        except Exception:
            return None

    try:
        ole = olefile.OleFileIO(file_path)
        if not ole.exists("WordDocument"):
            return None

        # 尝试读取文本流
        text_parts = []

        # 尝试 UTF-16 解码
        try:
            with ole.openstream("WordDocument") as f:
                doc_data = f.read()
            text_candidate = doc_data.decode("utf-16", errors="ignore")
            text_clean = "".join(c for c in text_candidate if c.isprintable() or c in "\n\r\t")
            if len(text_clean) > 100:
                text_parts.append(text_clean[:50000])
        except:
            pass

        # 尝试 ANSI (GBK) 解码
        try:
            with ole.openstream("WordDocument") as f:
                doc_data = f.read()
            text_candidate = doc_data.decode("gbk", errors="ignore")
            text_clean = "".join(c for c in text_candidate if c.isprintable() or c in "\n\r\t")
            if len(text_clean) > 100:
                text_parts.append(text_clean[:50000])
        except:
            pass

        if text_parts:
            return " ".join(text_parts)
        return None

    except Exception as e:
        print(f"  [ERROR] DOC 读取失败: {e}")
        return None


def read_txt(file_path: str) -> str:
    """读取 TXT 文件内容"""
    try:
        encodings = ["utf-8", "gbk", "gb2312", "gb18030"]
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        print(f"  [ERROR] TXT 读取失败: {e}")
        return None


def is_pdf(file_path: str) -> bool:
    """检测是否为 PDF 文件"""
    try:
        with open(file_path, "rb") as f:
            header = f.read(5)
            return header[:5] == b"%PDF-"
    except:
        return False


def scan_local_files() -> list:
    """扫描本地附件目录，返回附件列表"""
    attachments = []

    if not ATTACHMENTS_DIR.exists():
        print(f"[WARN] 附件目录不存在: {ATTACHMENTS_DIR}")
        return attachments

    for tender_id in os.listdir(ATTACHMENTS_DIR):
        tender_dir = ATTACHMENTS_DIR / tender_id
        if not tender_dir.is_dir():
            continue

        for filename in os.listdir(tender_dir):
            file_path = tender_dir / filename
            if file_path.is_file():
                attachments.append({
                    "tender_id": tender_id,
                    "file_name": filename,
                    "local_path": str(file_path),
                    "file_size": os.path.getsize(file_path)
                })

    print(f"[INFO] 扫描到 {len(attachments)} 个附件文件")
    return attachments


def process_attachments(tender_ids: list = None, limit: int = None):
    """
    处理附件并提取技术参数

    Args:
        tender_ids: 指定处理的 tender_id 列表，None 表示处理所有
        limit: 处理数量上限
    """
    if not HAS_OPENAI:
        print("[ERROR] openai 库未安装")
        return

    if not HAS_FITZ:
        print("[ERROR] PyMuPDF 未安装，PDF 处理不可用")
        return

    # 1. 获取附件列表
    print("=" * 60)
    print("扫描附件文件...")
    attachments = scan_local_files()

    if not attachments:
        print("[INFO] 没有找到附件文件")
        return

    # 过滤指定 tender_id
    if tender_ids:
        attachments = [a for a in attachments if str(a.get("tender_id")) in [str(t) for t in tender_ids]]

    if limit:
        attachments = attachments[:limit]

    print(f"待处理附件数: {len(attachments)}")
    print("=" * 60)

    # 2. 输出目录
    output_dir = BASE_DIR / "data" / "output" / "attachment_specs"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / "attachment_specs.jsonl"

    # 3. 逐个处理
    stats = {"total": len(attachments), "success": 0, "failed": 0, "skipped": 0}

    for idx, att in enumerate(attachments, 1):
        tender_id = str(att.get("tender_id", "unknown"))
        file_name = att.get("file_name", "")
        local_path = att.get("local_path", "")

        if not local_path or not os.path.exists(local_path):
            print(f"[{idx}/{len(attachments)}] [SKIP] 文件不存在: {local_path}")
            stats["skipped"] += 1
            continue

        print(f"\n[{idx}/{len(attachments)}] 处理中...")
        print(f"  文件: {file_name}")
        print(f"  路径: {local_path}")

        ext = os.path.splitext(file_name)[1].lower()

        # 根据文件类型选择处理方式
        if ext == ".pdf" and is_pdf(local_path):
            # PDF 文件: 使用视觉模型提取表格
            print("  [MODE] PDF 视觉提取模式")

            # 查找可能包含表格的页面
            table_pages = find_table_pages(local_path)
            print(f"  [INFO] 找到 {len(table_pages)} 页可能包含表格")

            # 转换前10页为图片
            images = pdf_to_images(local_path, start_page=0, end_page=10, dpi=150)
            if not images:
                print(f"  [FAIL] PDF 转图片失败")
                stats["failed"] += 1
                continue

            print(f"  [INFO] 转换了 {len(images)} 页为图片")

            # 调用视觉模型提取表格
            specs = extract_tables_with_vision(images, file_name, tender_id)

            if specs:
                stats["success"] += 1
                formatted_items = specs.get("formatted_items", [])
                item_count = len(formatted_items)
                print(f"  [OK] 提取到 {item_count} 个产品")

                # 每个 item 保存为单独一行
                for item in formatted_items:
                    record = {
                        "tender_id": tender_id,
                        "file_name": file_name,
                        **item  # 产品名称, 技术需求
                    }
                    save_jsonl([record], str(output_file), append=True)
            else:
                stats["failed"] += 1
                print(f"  [FAIL] 表格提取失败")

        elif ext in [".docx"]:
            # DOCX 文件: 读取文本后用文本模型
            print("  [MODE] DOCX 文本提取模式")
            text = read_docx(local_path)
            if not text or len(text.strip()) < 50:
                print(f"  [SKIP] 文本内容过少")
                stats["skipped"] += 1
                continue

            specs = extract_text_with_llm(text, file_name, tender_id)
            if specs:
                stats["success"] += 1
                save_jsonl([specs], str(output_file), append=True)
            else:
                stats["failed"] += 1

        elif ext in [".doc"]:
            # DOC 文件: 使用 olefile 读取文本后用文本模型
            print("  [MODE] DOC 文本提取模式")
            text = read_doc(local_path)
            if not text or len(text.strip()) < 50:
                print(f"  [SKIP] 文本内容过少或读取失败")
                stats["skipped"] += 1
                continue

            specs = extract_text_with_llm(text, file_name, tender_id)
            if specs:
                stats["success"] += 1
                save_jsonl([specs], str(output_file), append=True)
            else:
                stats["failed"] += 1

        elif ext in [".txt", ".csv"]:
            # TXT/CSV 文件: 直接用文本模型
            print("  [MODE] 文本文件提取模式")
            text = read_txt(local_path)
            if not text or len(text.strip()) < 50:
                print(f"  [SKIP] 文本内容过少")
                stats["skipped"] += 1
                continue

            specs = extract_text_with_llm(text, file_name, tender_id)
            if specs:
                stats["success"] += 1
                save_jsonl([specs], str(output_file), append=True)
            else:
                stats["failed"] += 1

        else:
            print(f"  [SKIP] 不支持的文件格式: {ext}")
            stats["skipped"] += 1

        if idx % 10 == 0:
            print(f"  [PROGRESS] 已处理 {idx}/{len(attachments)}")

    print("\n" + "=" * 60)
    print("处理完成!")
    print("=" * 60)
    print(f"总数: {stats['total']}")
    print(f"成功: {stats['success']}")
    print(f"失败: {stats['failed']}")
    print(f"跳过: {stats['skipped']}")
    print(f"输出: {output_file}")
    print("=" * 60)


def build_specs_index():
    """
    为提取的技术参数构建向量索引
    """
    from src.vectorization.vectorize_data import Vectorizer

    specs_file = BASE_DIR / "data" / "output" / "attachment_specs" / "attachment_specs.jsonl"

    if not specs_file.exists():
        print(f"[ERROR] 规格文件不存在: {specs_file}")
        return

    specs_data = load_jsonl(str(specs_file))
    if not specs_data:
        print("[ERROR] 没有规格数据")
        return

    print(f"为 {len(specs_data)} 条规格数据构建索引...")

    # 构建文本用于向量化
    texts = []
    for spec in specs_data:
        # 对于表格数据，提取产品名称和技术参数
        if spec.get("source_type") == "pdf_vision":
            products = spec.get("products", [])
            text_parts = []
            for p in products:
                name = p.get("产品名称", "")
                spec_text = p.get("规格型号", "")
                tech = p.get("技术参数", "")
                if name:
                    text_parts.append(f"{name} {spec_text} {tech}")
            text = " | ".join(text_parts)
        else:
            # 文本模式
            text_parts = [
                " ".join(spec.get("product_names", [])),
                " ".join(spec.get("specifications", [])),
                " ".join(spec.get("technical_params", [])),
            ]
            text = " | ".join(filter(None, text_parts))
        texts.append(text)

    # 向量化
    vectorizer = Vectorizer()
    embeddings = vectorizer.get_embeddings(texts)

    # 保存索引
    import faiss
    import numpy as np

    dimension = len(embeddings[0])
    index = faiss.IndexFlatIP(dimension)

    vecs = np.array(embeddings).astype("float32")
    faiss.normalize_L2(vecs)
    index.add(vecs)

    # 保存
    index_dir = BASE_DIR / "data" / "index_specs"
    index_dir.mkdir(parents=True, exist_ok=True)
    faiss.write_index(index, str(index_dir / "specs.index"))

    # 保存 ID 映射
    ids_file = index_dir / "specs_ids.json"
    with open(ids_file, "w", encoding="utf-8") as f:
        json.dump([{
            "tender_id": s.get("tender_id"),
            "file_name": s.get("file_name"),
            "product_count": len(s.get("products", [])) if s.get("source_type") == "pdf_vision" else 0
        } for s in specs_data], f, ensure_ascii=False)

    print(f"[OK] 索引已保存到 {index_dir}")
    print(f"  - specs.index: {index.ntotal} 条向量")
    print(f"  - specs_ids.json: ID 映射")


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="附件技术参数提取")
    parser.add_argument("--tender-ids", nargs="*", help="指定处理的 tender_id")
    parser.add_argument("--limit", type=int, default=None, help="处理数量上限")
    parser.add_argument("--build-index", action="store_true", help="构建向量索引")

    args = parser.parse_args()

    if args.build_index:
        build_specs_index()
    else:
        process_attachments(tender_ids=args.tender_ids, limit=args.limit)
